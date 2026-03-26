from docx import Document
from zipfile import ZipFile
import xml.etree.ElementTree as ET

path = r'c:\Users\31897\Desktop\作业、实验报告批改案例\面向对象程序设计（Java）（作业）\对象与类作业-批改 - 副本.docx'
doc = Document(path)

print('=== PARAGRAPHS ===')
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f'[{i}] Style: {p.style.name} | Text: {repr(p.text)}')

print()
print('=== TABLES ===')
for t_idx, table in enumerate(doc.tables):
    print(f'Table {t_idx}:')
    for row in table.rows:
        for cell in row.cells:
            if cell.text.strip():
                print(f'  Cell: {repr(cell.text)}')

# Read comments from the docx package
try:
    with ZipFile(path, 'r') as z:
        names = z.namelist()
        print()
        print('=== ZIP CONTENTS (word/) ===')
        for n in names:
            if 'word/' in n:
                print(f'  {n}')

        if 'word/comments.xml' in names:
            comments_xml = z.read('word/comments.xml').decode('utf-8')
            print()
            print('=== COMMENTS XML (first 5000 chars) ===')
            print(comments_xml[:5000])
        else:
            print('No word/comments.xml found')
except Exception as e:
    print(f'Error: {e}')
