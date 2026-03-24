"""
导入 rag-data 目录中的文件到 RAG 系统

使用方法:
    python import_rag_data.py

说明:
    - 自动扫描 rag-data 目录下的所有文件
    - 按课程目录组织 (course_id)
    - 直接写入数据库和索引 (不依赖 API)
    - 支持扫描版 PDF 的 OCR 识别
    - 支持 DOC 文件的正确解析
"""

import asyncio
import os
import sys
import hashlib
import io
import re
import tempfile
import subprocess
from pathlib import Path
from typing import List, Tuple, Dict, Any, Optional
from datetime import datetime
import uuid

# 添加 backend 到路径
sys.path.insert(0, str(Path(__file__).parent))

# 配置
RAG_DATA_DIR = Path(__file__).parent.parent / "rag-data"

# 课程 ID 映射 (根据目录名/文件名自动生成)
def get_course_id_from_dir(dir_name: str) -> str:
    """从目录名生成 course_id"""
    return dir_name.replace("（", "(").replace("）", ")").replace(" ", "_")


def get_course_id_from_filename(filename: str) -> str:
    """从文件名生成 course_id"""
    # 移除扩展名
    name = Path(filename).stem
    
    # 根据常见教材名称映射
    if "java" in name.lower():
        if "core" in name.lower() or "fundamentals" in name.lower():
            return "Java 核心技术"
        return "Java 编程"
    elif "python" in name.lower():
        return "Python 编程"
    elif "experiment" in name.lower() or "实验" in name:
        return "实验报告"
    elif "homework" in name.lower() or "作业" in name:
        return "作业练习"
    else:
        # 使用文件名前缀作为 course_id
        return name[:20].replace(" ", "_").replace("-", "_")


def extract_text_from_pdf_plumber(file_content: bytes) -> str:
    """从 PDF 提取文本 (使用 pdfplumber - 适用于文本型 PDF)"""
    import pdfplumber
    text_parts = []
    try:
        with pdfplumber.open(io.BytesIO(file_content)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text and text.strip():
                    text_parts.append(text)
    except Exception as e:
        print(f"[pdfplumber] 提取失败：{e}")
        return ""
    
    if text_parts:
        return '\n\n'.join(text_parts)
    return ""


def extract_text_from_pdf_pymupdf(file_content: bytes) -> str:
    """从 PDF 提取文本 (使用 PyMuPDF/fitz - 更强的文本提取能力)"""
    try:
        import fitz  # PyMuPDF
        text_parts = []
        doc = fitz.open(stream=io.BytesIO(file_content), filetype="pdf")
        for page in doc:
            text = page.get_text("text")
            if text and text.strip():
                text_parts.append(text)
        doc.close()
        
        if text_parts:
            return '\n\n'.join(text_parts)
    except Exception as e:
        print(f"[PyMuPDF] 提取失败：{e}")
    
    return ""


def extract_text_from_pdf_ocr(file_content: bytes, use_llm: bool = True) -> str:
    """
    从 PDF 提取文本 (使用 OCR - 适用于扫描版 PDF)
    
    方案 1: 使用 pymupdf4llm (推荐，支持 Markdown 格式输出)
    方案 2: 使用 pdf2image + pytesseract (传统 OCR)
    """
    text_parts = []
    
    # 方案 1: 尝试使用 pymupdf4llm (更好的 OCR 支持)
    if use_llm:
        try:
            import fitz  # PyMuPDF
            import pymupdf4llm
            
            # 保存临时文件
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
                tmp.write(file_content)
                tmp_path = tmp.name
            
            try:
                # 使用 pymupdf4llm 提取 (支持 OCR)
                md_text = pymupdf4llm.to_markdown(tmp_path)
                if md_text and md_text.strip():
                    text_parts.append(md_text)
                    print("[OCR-pymupdf4llm] 成功提取")
            finally:
                os.unlink(tmp_path)
                
        except ImportError:
            print("[OCR-pymupdf4llm] 未安装，跳过")
        except Exception as e:
            print(f"[OCR-pymupdf4llm] 提取失败：{e}")
    
    # 方案 2: 使用 pdf2image + pytesseract
    if not text_parts:
        try:
            from pdf2image import convert_from_bytes
            import pytesseract
            
            # 转换 PDF 为图片
            images = convert_from_bytes(file_content, dpi=300)
            print(f"[OCR-tesseract] 转换了 {len(images)} 页")
            
            for i, image in enumerate(images):
                # OCR 识别
                text = pytesseract.image_to_string(image, lang='chi_sim+eng')
                if text and text.strip():
                    text_parts.append(f"--- 第 {i+1} 页 ---\n{text}")
            
            if text_parts:
                print("[OCR-tesseract] 成功提取")
                
        except ImportError:
            print("[OCR-tesseract] 未安装 pdf2image 或 pytesseract，跳过")
        except Exception as e:
            print(f"[OCR-tesseract] 提取失败：{e}")
    
    return '\n\n'.join(text_parts) if text_parts else ""


def extract_text_from_pdf(file_content: bytes) -> str:
    """
    从 PDF 提取文本 (多策略)
    
    策略顺序:
    1. pdfplumber (快速，适用于文本型 PDF)
    2. PyMuPDF (更强的文本提取)
    3. OCR (适用于扫描版 PDF)
    """
    print("\n    [PDF] 开始提取文本...")
    
    # 策略 1: pdfplumber
    text = extract_text_from_pdf_plumber(file_content)
    if text and len(text.strip()) > 100:
        print(f"    [PDF] 使用 pdfplumber 提取了 {len(text)} 字符")
        return text
    
    # 策略 2: PyMuPDF
    text = extract_text_from_pdf_pymupdf(file_content)
    if text and len(text.strip()) > 100:
        print(f"    [PDF] 使用 PyMuPDF 提取了 {len(text)} 字符")
        return text
    
    # 策略 3: OCR (扫描版)
    print("    [PDF] 文本提取不足，尝试 OCR 识别...")
    text = extract_text_from_pdf_ocr(file_content)
    if text and len(text.strip()) > 100:
        print(f"    [PDF] 使用 OCR 提取了 {len(text)} 字符")
        return text
    
    # 所有策略都失败
    print("    [PDF] 所有提取方法都失败")
    return ""


def extract_text_from_docx(file_content: bytes) -> str:
    """从 DOCX 提取文本"""
    from docx import Document
    try:
        doc = Document(io.BytesIO(file_content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # 也提取表格内容
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                if any(row_text):
                    tables_text.append(" | ".join(row_text))
        
        all_text = paragraphs + tables_text
        return '\n\n'.join(all_text) if all_text else ""
    except Exception as e:
        print(f"[DOCX] 提取失败：{e}")
        return ""


def extract_text_from_doc_olefile(file_content: bytes) -> str:
    """
    从 DOC 文件提取文本 (使用 olefile + 字符串提取)
    适用于老版本 Word 文档 (.doc)
    """
    try:
        import olefile
        
        ole = olefile.OleFileIO(io.BytesIO(file_content))
        text_parts = []
        
        # 尝试从不同的流中提取文本
        # WordDocument 流包含文档的文本内容
        if ole.exists("WordDocument"):
            doc_stream = ole.openstream("WordDocument")
            doc_data = doc_stream.read()
            # 尝试提取可打印的文本
            text = extract_text_from_word_document_stream(doc_data)
            if text:
                text_parts.append(text)
        
        # 尝试从 0Table 或 1Table 流中提取 (包含格式化信息)
        for table_stream in ["0Table", "1Table"]:
            if ole.exists(table_stream):
                try:
                    stream = ole.openstream(table_stream)
                    data = stream.read()
                    # 尝试提取 Unicode 文本
                    text = extract_unicode_text_from_stream(data)
                    if text:
                        text_parts.append(f"[{table_stream}] {text}")
                except:
                    pass
        
        ole.close()
        
        if text_parts:
            return '\n\n'.join(text_parts)
            
    except ImportError:
        print("[DOC-olefile] olefile 未安装，跳过")
    except Exception as e:
        print(f"[DOC-olefile] 提取失败：{e}")
    
    return ""


def extract_text_from_word_document_stream(data: bytes) -> str:
    """从 WordDocument 流中提取文本"""
    text_parts = []
    
    # WordDocument 流的结构复杂，这里尝试简单提取
    # 跳过文件头 (通常是 0x200 字节)
    offset = 0x200
    
    # 查找并提取 FIB (File Information Block)
    if len(data) < offset + 4:
        return ""
    
    # 尝试提取文本片段
    # 这是一种简化的方法，实际 DOC 格式更复杂
    current_pos = offset
    while current_pos < len(data) - 1:
        # 查找可能的文本开始位置
        if data[current_pos:current_pos+2] != b'\x00\x00':
            # 尝试读取一段数据
            chunk_end = min(current_pos + 1024, len(data))
            chunk = data[current_pos:chunk_end]
            
            # 提取可打印字符
            text_chunk = ""
            for b in chunk:
                if 32 <= b <= 126 or b in [9, 10, 13]:  # 可打印 ASCII + 空白字符
                    text_chunk += chr(b)
                elif b == 0:  # 遇到 null 终止符
                    break
            
            if len(text_chunk.strip()) > 10:
                text_parts.append(text_chunk.strip())
            
            # 跳到下一个可能的文本块
            current_pos = chunk_end + 100  # 跳过一些字节
        else:
            current_pos += 1
    
    return '\n'.join(text_parts) if text_parts else ""


def extract_unicode_text_from_stream(data: bytes) -> str:
    """从二进制流中提取 Unicode 文本"""
    text_parts = []
    
    # 尝试提取 UTF-16LE 编码的文本 (Windows 常用)
    try:
        # 查找连续的 Unicode 可打印字符
        i = 0
        while i < len(data) - 1:
            # 跳过非文本区域
            if data[i:i+2] == b'\x00\x00':
                i += 2
                continue
            
            # 尝试读取 UTF-16LE 字符
            char_bytes = data[i:i+2]
            if len(char_bytes) == 2:
                try:
                    char = char_bytes.decode('utf-16-le')
                    if char.isprintable() or char in ['\n', '\r', '\t']:
                        text_parts.append(char)
                    else:
                        # 遇到非可打印字符，检查是否收集了足够的文本
                        if len(text_parts) > 10:
                            break
                except:
                    pass
            i += 2
        
        result = ''.join(text_parts).strip()
        if len(result) > 10:
            return result
    except:
        pass
    
    return ""


def extract_text_from_doc_antiword(file_path: Path) -> str:
    """使用 antiword 提取 DOC 文件文本"""
    try:
        result = subprocess.run(
            ['antiword', '-t', str(file_path)],
            capture_output=True,
            text=True,
            timeout=30
        )
        if result.returncode == 0 and result.stdout.strip():
            print("[DOC-antiword] 成功提取")
            return result.stdout
    except subprocess.TimeoutExpired:
        print("[DOC-antiword] 超时")
    except FileNotFoundError:
        print("[DOC-antiword] antiword 未安装，请运行：sudo apt-get install antiword")
    except Exception as e:
        print(f"[DOC-antiword] 提取失败：{e}")
    
    return ""


def extract_text_from_doc(file_path: Path, file_content: bytes) -> str:
    """
    从 DOC 文件提取文本 (多策略)
    
    策略顺序:
    1. antiword (专门工具)
    2. olefile + 手动解析
    3. 降级处理 (二进制读取)
    """
    print("\n    [DOC] 开始提取文本...")
    
    # 策略 1: antiword
    text = extract_text_from_doc_antiword(file_path)
    if text and len(text.strip()) > 100:
        print(f"    [DOC] 使用 antiword 提取了 {len(text)} 字符")
        return text
    
    # 策略 2: olefile
    text = extract_text_from_doc_olefile(file_content)
    if text and len(text.strip()) > 100:
        print(f"    [DOC] 使用 olefile 提取了 {len(text)} 字符")
        return text
    
    # 策略 3: 降级处理 - 尝试提取二进制中的文本
    print("    [DOC] 尝试降级提取...")
    text = extract_fallback_text_from_binary(file_content)
    if text and len(text.strip()) > 100:
        print(f"    [DOC] 降级提取了 {len(text)} 字符")
        return text
    
    print("    [DOC] 所有提取方法都失败")
    return ""


def extract_fallback_text_from_binary(file_content: bytes) -> str:
    """降级方案：从二进制文件中提取可识别的文本"""
    text_parts = []
    current_text = []
    
    for byte in file_content:
        # 可打印 ASCII 字符 + 常见空白
        if 32 <= byte <= 126 or byte in [9, 10, 13]:
            current_text.append(chr(byte))
        else:
            # 遇到非文本字符，保存之前收集的文本
            if len(current_text) > 20:  # 至少 20 个字符才保留
                text_parts.append(''.join(current_text).strip())
            current_text = []
    
    # 处理最后一段
    if len(current_text) > 20:
        text_parts.append(''.join(current_text).strip())
    
    # 过滤掉明显的二进制垃圾
    filtered_parts = []
    for part in text_parts:
        # 检查是否包含合理的单词/字符比例
        words = part.split()
        if len(words) >= 3:  # 至少有 3 个"词"
            filtered_parts.append(part)
    
    return '\n\n'.join(filtered_parts) if filtered_parts else ""


def extract_text_from_file(file_path: Path) -> str:
    """从文件提取文本"""
    ext = file_path.suffix.lower()

    if ext == '.pdf':
        with open(file_path, 'rb') as f:
            return extract_text_from_pdf(f.read())
    elif ext in ['.docx']:
        with open(file_path, 'rb') as f:
            return extract_text_from_docx(f.read())
    elif ext in ['.doc']:
        with open(file_path, 'rb') as f:
            content = f.read()
        return extract_text_from_doc(file_path, content)
    elif ext in ['.java', '.txt', '.md', '.py', '.js', '.ts']:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    else:
        # 未知类型，尝试读取为文本
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()


def scan_rag_data_directory() -> List[Tuple[Path, str]]:
    """扫描 rag-data 目录，返回 (文件路径，course_id) 列表"""
    files_to_import = []

    if not RAG_DATA_DIR.exists():
        print(f"错误：rag-data 目录不存在：{RAG_DATA_DIR}")
        return files_to_import

    print(f"扫描目录：{RAG_DATA_DIR}")

    # 遍历目录
    for item in RAG_DATA_DIR.iterdir():
        if item.is_file():
            # 根目录的文件 - 根据文件名生成 course_id
            if item.suffix.lower() in ['.pdf', '.txt', '.md', '.doc', '.docx']:
                course_id = get_course_id_from_filename(item.name)
                files_to_import.append((item, course_id))
                print(f"  发现文件：{item.name} -> course_id={course_id}")

        elif item.is_dir():
            # 子目录 - 使用目录名作为 course_id
            course_id = get_course_id_from_dir(item.name)
            print(f"  发现课程目录：{item.name} -> course_id={course_id}")

            # 遍历子目录中的文件
            for file_item in item.rglob("*"):
                if file_item.is_file() and file_item.suffix.lower() in ['.pdf', '.txt', '.md', '.doc', '.docx', '.java']:
                    files_to_import.append((file_item, course_id))
                    print(f"    - {file_item.name}")

    return files_to_import


async def import_to_database(file_path: Path, course_id: str, content: str) -> Dict[str, Any]:
    """导入文件内容到数据库"""
    from common.database.postgresql import AsyncSessionLocal
    from services.knowledge.rag.main import RAGDocument
    from sqlalchemy import select
    
    # 分块
    from services.knowledge.chunk.main import chunk_text, ChunkRequest
    
    chunk_req = ChunkRequest(
        content=content, 
        content_type="教材", 
        doc_metadata={"filename": file_path.name, "course_id": course_id}
    )
    chunk_response = await chunk_text(chunk_req)
    chunks = chunk_response.chunks
    
    doc_ids = []
    
    async with AsyncSessionLocal() as session:
        for i, chunk in enumerate(chunks):
            doc_id = hashlib.md5(f"{file_path.name}_{i}_{course_id}".encode()).hexdigest()
            doc_ids.append(doc_id)
            
            # 检查是否已存在
            existing = await session.execute(
                select(RAGDocument).where(RAGDocument.doc_id == doc_id)
            )
            if not existing.scalar_one_or_none():
                rag_doc = RAGDocument(
                    doc_id=doc_id,
                    content=chunk["content"],
                    doc_metadata={
                        "filename": file_path.name,
                        "chunk_index": i,
                        "course_id": course_id
                    },
                    course_id=course_id
                )
                session.add(rag_doc)
        
        await session.commit()
    
    return {
        "success": True,
        "filename": file_path.name,
        "course_id": course_id,
        "doc_count": len(doc_ids),
        "doc_ids": doc_ids
    }


async def build_indices():
    """构建 TF-IDF 和 BM25 索引（使用 TF-IDF 嵌入构建 FAISS 索引）"""
    from services.knowledge.index_builder.main import build_tfidf_index_from_db, build_faiss_index_from_db
    from services.knowledge.bm25_search.main import bm25_index, BM25Document
    from common.database.postgresql import AsyncSessionLocal
    from sqlalchemy import select
    from services.knowledge.rag.main import RAGDocument

    tfidf_count = 0
    bm25_count = 0
    faiss_count = 0

    print("构建 TF-IDF 索引...")
    try:
        tfidf_count = await build_tfidf_index_from_db()
        print(f"  TF-IDF 索引构建完成：{tfidf_count} 个文档")
    except Exception as e:
        print(f"  TF-IDF 索引构建失败：{e}")

    print("构建 FAISS 索引（使用 TF-IDF 嵌入）...")
    try:
        faiss_count = await build_faiss_index_from_db()
        print(f"  FAISS 索引构建完成：{faiss_count} 个文档")
    except Exception as e:
        print(f"  FAISS 索引构建失败：{e}")

    print("构建 BM25 索引...")
    # 从数据库加载所有文档并添加到 BM25 索引
    # 注意：必须创建新的 BM25Index 实例，不能使用全局的 bm25_index
    # 因为全局实例可能已有旧数据，导致 save() 时数据不一致
    try:
        from services.knowledge.bm25_search.main import BM25Index as NewBM25Index
        
        # 创建新的索引实例
        local_bm25_index = NewBM25Index()
        
        async with AsyncSessionLocal() as session:
            result = await session.execute(select(RAGDocument))
            docs = result.scalars().all()
            print(f"  从数据库加载了 {len(docs)} 个文档")

            for doc in docs:
                bm25_doc = BM25Document(
                    doc_id=doc.doc_id,
                    content=doc.content,
                    course_id=doc.course_id,
                    metadata=doc.doc_metadata or {}
                )
                local_bm25_index.add_document(bm25_doc)

        # 保存索引到磁盘
        local_bm25_index.save()
        bm25_count = len(local_bm25_index.documents)
        print(f"  BM25 索引构建完成：{bm25_count} 个文档")
        
        # 同步到全局实例（用于当前会话）
        bm25_index.documents = local_bm25_index.documents
        bm25_index.doc_lengths = local_bm25_index.doc_lengths
        bm25_index.avg_doc_length = local_bm25_index.avg_doc_length
        bm25_index.inverted_index = local_bm25_index.inverted_index
        bm25_index.doc_frequency = local_bm25_index.doc_frequency
        bm25_index.total_docs = local_bm25_index.total_docs
        bm25_index.vocabulary = local_bm25_index.vocabulary
        
    except Exception as e:
        print(f"  BM25 索引构建失败：{e}")
        import traceback
        traceback.print_exc()
        # 注意：不要保存空索引，以免覆盖已有的有效索引
        bm25_count = 0

    return faiss_count, tfidf_count, bm25_count


async def main():
    """主函数"""
    print("=" * 60)
    print("RAG 数据导入工具 (直接导入模式)")
    print("=" * 60)

    # 扫描文件
    print("\n[步骤 1/3] 扫描 rag-data 目录...")
    files_to_import = scan_rag_data_directory()

    if not files_to_import:
        print("未发现需要导入的文件")
        return

    print(f"\n共发现 {len(files_to_import)} 个文件")

    # 提取文本并导入数据库
    print(f"\n[步骤 2/3] 提取文本并导入数据库...")
    print("-" * 60)

    results = {"success": 0, "failed": 0, "total_chunks": 0}
    import_results = []

    for file_path, course_id in files_to_import:
        print(f"处理：{file_path.name} (course_id={course_id})...", end=" ")
        
        try:
            # 提取文本
            content = extract_text_from_file(file_path)
            
            if not content or len(content.strip()) < 10:
                print(f"跳过：文本内容太少")
                results["failed"] += 1
                import_results.append({
                    "success": False,
                    "filename": file_path.name,
                    "message": "文本内容太少"
                })
                continue
            
            print(f"提取了 {len(content)} 字符，导入数据库...", end=" ")
            
            # 导入数据库
            result = await import_to_database(file_path, course_id, content)
            import_results.append(result)
            
            if result["success"]:
                print(f"成功 (分块数：{result['doc_count']})")
                results["success"] += 1
                results["total_chunks"] += result["doc_count"]
            else:
                print(f"失败：{result.get('message', '未知错误')}")
                results["failed"] += 1
                
        except Exception as e:
            print(f"失败：{str(e)[:50]}")
            results["failed"] += 1
            import_results.append({
                "success": False,
                "filename": file_path.name,
                "message": str(e)
            })

    print("-" * 60)
    print(f"导入完成：成功 {results['success']} 个，失败 {results['failed']} 个，共 {results['total_chunks']} 个文档块")

    # 构建索引
    print(f"\n[步骤 3/3] 构建离线索引...")
    try:
        faiss_count, tfidf_count, bm25_count = await build_indices()
        print(f"索引构建完成!")
        print(f"  - FAISS 向量数：{faiss_count} (使用 TF-IDF 嵌入)")
        print(f"  - TF-IDF 向量数：{tfidf_count}")
        print(f"  - BM25 文档数：{bm25_count}")
    except Exception as e:
        print(f"索引构建失败：{e}")
        import traceback
        traceback.print_exc()
        print("提示：请确保后端服务启动后运行 build_index 接口")

    # 总结
    print("\n" + "=" * 60)
    print("导入完成!")
    print("=" * 60)

    if results["failed"] > 0:
        print("\n失败的文件:")
        for r in import_results:
            if not r.get("success"):
                print(f"  - {r.get('filename', 'N/A')}: {r.get('message', '未知错误')}")


if __name__ == "__main__":
    asyncio.run(main())
