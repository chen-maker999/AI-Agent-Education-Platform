"""作业批改服务 - AI自动批改作业并添加docx批注"""

import os
import re
import uuid
import hashlib
import io
import zipfile
import json
from datetime import datetime, timedelta
from typing import Optional, List
from fastapi import APIRouter, UploadFile, File, Form, HTTPException, Query
from pydantic import BaseModel, Field
from sqlalchemy import Column, String, Integer, DateTime, Text, Float
from sqlalchemy.dialects.postgresql import UUID, JSONB
from common.models.response import ResponseModel
from common.core.config import settings
from common.database.postgresql import Base, AsyncSessionLocal
from docx.oxml import OxmlElement
from minio import Minio
from minio.error import S3Error
from services.data.homework.main import Homework as Homework

router = APIRouter(prefix="/homework_review", tags=["Homework Review"])

BUCKET_REVIEW = "homework-reviews"


# SQLAlchemy Model - matches existing DB schema
class HomeworkReview(Base):
    __tablename__ = "homework_reviews"

    id = Column(UUID(as_uuid=True), primary_key=True, default=uuid.uuid4)
    review_id = Column("review_id", String(255), unique=True, index=True, nullable=True)
    homework_id = Column("homework_id", String(255), nullable=False, index=True)
    student_id = Column("student_id", String(255), nullable=False, index=True)
    original_filename = Column("original_filename", String(500), nullable=True)
    graded_filename = Column("graded_filename", String(500), nullable=True)
    graded_file_url = Column("graded_file_url", Text, nullable=True)
    file_size = Column("file_size", Integer, nullable=True)
    score = Column("score", Float, nullable=True)
    total_score = Column("max_score", Float, nullable=True)  # DB uses max_score
    status = Column("grading_status", String(50), default="pending")  # DB uses grading_status
    course = Column("course", String(100), nullable=True)
    issue_count = Column("issue_count", Integer, default=0)
    issues_json = Column("issues_json", Text, nullable=True)
    code_issues = Column("code_issues", JSONB, nullable=True)
    original_content = Column("original_content", Text, nullable=True)
    grading_model = Column("grading_model", String(100), nullable=True)
    grading_time = Column("grading_time", DateTime, nullable=True)
    review_details = Column("review_details", JSONB, nullable=True)
    created_at = Column("created_at", DateTime, default=datetime.utcnow)
    updated_at = Column("updated_at", DateTime, default=datetime.utcnow)


def get_minio_client() -> Minio:
    return Minio(
        settings.MINIO_ENDPOINT,
        access_key=settings.MINIO_ACCESS_KEY,
        secret_key=settings.MINIO_SECRET_KEY,
        secure=settings.MINIO_SECURE
    )


def ensure_bucket():
    try:
        client = get_minio_client()
        if not client.bucket_exists(BUCKET_REVIEW):
            client.make_bucket(BUCKET_REVIEW)
    except Exception:
        pass


def read_docx_text(content: bytes) -> str:
    """从docx/doc内容中提取文本，支持：
    - .docx (Open XML / ZIP格式)
    - .doc (OLE2 / Microsoft Office 97-2003 格式) - 仅用于文本提取
    """
    import logging
    logger = logging.getLogger(__name__)

    # 检查内容是否为空
    if not content or len(content) < 4:
        logger.error(f"文件内容为空或太短: {len(content) if content else 0} bytes")
        return ""

    # 检查文件格式
    file_header = content[:4].hex()
    logger.info(f"文件头: {file_header}")

    # 方法1: docx 格式 (ZIP with PK header)
    if content[:2] == b'PK':
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
            parts = []
            for p in doc.paragraphs:
                if p.text.strip():
                    parts.append(p.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            parts.append(cell.text)
            text = '\n'.join(parts)
            if text.strip():
                logger.info(f"成功从 docx 提取 {len(text)} 字符")
                return text
        except Exception as e:
            logger.warning(f"python-docx 读取失败: {e}")

        try:
            with zipfile.ZipFile(io.BytesIO(content), 'r') as z:
                if 'word/document.xml' in z.namelist():
                    with z.open('word/document.xml') as f:
                        xml = f.read().decode('utf-8')
                        text = re.sub(r'<[^>]+>', ' ', xml)
                        text = re.sub(r'\s+', '\n', text).strip()
                        logger.info(f"从 XML 提取 {len(text)} 字符")
                        return text
        except Exception as e:
            logger.error(f"ZIP/XML 读取失败: {e}")

    # 方法2: doc 格式 (OLE2 / CFEF header)
    # d0cf11e0 是 OLE2 文件签名
    if content[:4] == b'\xd0\xcf\x11\xe0':
        logger.error("检测到 .doc 格式（旧版 Word 97-2003 格式），需要转换为 .docx 格式才能添加批注")

        # 尝试提取文本用于评分，但不返回（因为无法写入批注）
        try:
            import olefile
            ole = olefile.OleFileIO(io.BytesIO(content))

            text_parts = []

            try:
                word_stream = ole.openstream('WordDocument')
                word_data = word_stream.read()

                current_line = []
                for i in range(512, len(word_data)):
                    byte = word_data[i]
                    if 32 <= byte <= 126:
                        current_line.append(chr(byte))
                    elif byte in (9, 10, 13):
                        current_line.append(' ')
                    else:
                        if len(current_line) >= 3:
                            line = ''.join(current_line).strip()
                            if line and not line.startswith('\x00'):
                                text_parts.append(line)
                        current_line = []

                for table_name in ['1Table', '0Table', 'WordDocument']:
                    if ole.exists(table_name):
                        try:
                            table_stream = ole.openstream(table_name)
                            table_data = table_stream.read()
                            current_line = []
                            for i in range(0, min(len(table_data), 50000)):
                                byte = table_data[i]
                                if 32 <= byte <= 126:
                                    current_line.append(chr(byte))
                                elif byte in (9, 10, 13):
                                    current_line.append(' ')
                                else:
                                    if len(current_line) >= 3:
                                        line = ''.join(current_line).strip()
                                        if line and not line.startswith('\x00'):
                                            text_parts.append(line)
                                    current_line = []
                        except:
                            pass

            except Exception as e:
                logger.warning(f"读取 WordDocument stream 失败: {e}")

            text = ' '.join(text_parts)
            text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', text)
            text = re.sub(r'\s+', ' ', text).strip()

            logger.info(f"OLE2 提取文本长度: {len(text)}")

            if len(text) > 50:
                logger.info(f"成功从 doc (OLE2) 提取 {len(text)} 字符")

                # 返回特殊标记，表示这是 doc 格式
                return f"[DOC_FORMAT]{text}"

            ole.close()

        except ImportError:
            logger.warning("olefile 库未安装，无法读取 .doc 格式")
        except Exception as e:
            logger.error(f"OLE2/doc 读取失败: {e}")

    logger.error(f"无法识别文件格式，文件头: {file_header}")
    return ""


def grade_java_code(code_text: str) -> dict:
    """对Java代码作业进行评分，返回评分结果"""
    score = 100.0
    issues_data = []

    # 1. 检查类定义
    required_classes = ["EngineConfig", "Spaceship", "CargoShuttle", "ExplorerProbe"]
    for cls in required_classes:
        if f"class {cls}" not in code_text:
            issues_data.append({
                "type": "missing_class",
                "severity": "error",
                "class": cls,
                "description": f"缺少 {cls} 类定义",
                "score_impact": -15.0,
                "suggestion": f"请定义 class {cls} {{ }}",
                "location": None,
            })
            score -= 15.0

    # 2. 检查 EngineConfig 构造函数
    if "class EngineConfig" in code_text:
        ec_pattern = r'public\s+EngineConfig\s*\([^)]*engineID[^)]*\)'
        if not re.search(ec_pattern, code_text):
            issues_data.append({
                "type": "constructor",
                "severity": "error",
                "class": "EngineConfig",
                "description": "EngineConfig 构造函数参数未正确赋值给字段",
                "score_impact": -5.0,
                "suggestion": "请在构造函数中使用 this.engineID = engineID 进行赋值",
                "location": "EngineConfig 构造函数",
            })
            score -= 5.0

        if re.search(r'engineID\s*=\s*engineID\s*;', code_text):
            issues_data.append({
                "type": "this_keyword",
                "severity": "error",
                "class": "EngineConfig",
                "description": "构造函数中出现 engineID = engineID 自我赋值错误",
                "score_impact": -5.0,
                "suggestion": "请使用 this.engineID = engineID 以区分参数和字段",
                "location": "EngineConfig 构造函数",
            })
            score -= 5.0

    # 3. 检查 Spaceship.getTotalInstances
    if "class Spaceship" in code_text:
        if re.search(r'public\s+static\s+\w+\s+getTotalInstances\s*\(\s*\)', code_text):
            if re.search(r'\bthis\s*\.\s*totalInstances', code_text):
                issues_data.append({
                    "type": "static_this",
                    "severity": "error",
                    "class": "Spaceship",
                    "description": "静态方法中不能使用 this 关键字访问 totalInstances",
                    "score_impact": -5.0,
                    "suggestion": "请使用 Spaceship.totalInstances 或直接使用 totalInstances（静态方法可直接访问静态字段）",
                    "location": "Spaceship.getTotalInstances()",
                })
                score -= 5.0

        if re.search(r'public\s+\w+\s+clone\s*\(\s*\)', code_text):
            if 'cloned.engine' not in code_text and 'this.engine.clone()' not in code_text:
                issues_data.append({
                    "type": "shallow_clone",
                    "severity": "warning",
                    "class": "Spaceship",
                    "description": "Spaceship.clone() 可能存在浅克隆问题，engine对象未被深拷贝",
                    "score_impact": -10.0,
                    "suggestion": "请在clone()中添加 cloned.engine = this.engine.clone() 实现深克隆",
                    "location": "Spaceship.clone()",
                })
                score -= 10.0

    # 4. 检查 CargoShuttle
    if "class CargoShuttle" in code_text:
        if re.search(r'Spaceship\s*\([^)]*\)\s*;', code_text):
            issues_data.append({
                "type": "super_call",
                "severity": "error",
                "class": "CargoShuttle",
                "description": "构造函数中直接调用 Spaceship(...) 而非 super(...)",
                "score_impact": -5.0,
                "suggestion": "请使用 super(fuel, engineID); 调用父类构造函数",
                "location": "CargoShuttle 构造函数",
            })
            score -= 5.0

        if re.search(r'public\s+\w+\s+CargoShuttle\s+clone\s*\(\s*\)', code_text):
            if not re.search(r'cloned\s*\.\s*cargoCapacity\s*=', code_text):
                issues_data.append({
                    "type": "clone_field",
                    "severity": "warning",
                    "class": "CargoShuttle",
                    "description": "CargoShuttle.clone() 可能未拷贝 cargoCapacity 字段",
                    "score_impact": -5.0,
                    "suggestion": "请在clone()中添加 cloned.cargoCapacity = this.cargoCapacity",
                    "location": "CargoShuttle.clone()",
                })
                score -= 5.0

        if re.search(r'getFuel\s*\(\s*\)\s*\+\s*getEngine\s*\(\s*\)', code_text):
            issues_data.append({
                "type": "toString_syntax",
                "severity": "error",
                "class": "CargoShuttle",
                "description": "toString() 方法中缺少操作符，可能写成 getFuel() + getEngine()",
                "score_impact": -3.0,
                "suggestion": "请使用 getFuel() + \" \" + getEngine() 或格式化字符串",
                "location": "CargoShuttle.toString()",
            })
            score -= 3.0

    # 5. 检查 ExplorerProbe
    if "class ExplorerProbe" in code_text:
        if 'ArrayList<String> scientificInstruments' in code_text:
            if not re.search(r'new\s+ArrayList', code_text) or \
               'this.scientificInstruments = new ArrayList' not in code_text:
                issues_data.append({
                    "type": "null_pointer",
                    "severity": "error",
                    "class": "ExplorerProbe",
                    "description": "scientificInstruments 未在构造函数中初始化，会导致 NullPointerException",
                    "score_impact": -5.0,
                    "suggestion": "请在构造函数中添加 this.scientificInstruments = new ArrayList<>();",
                    "location": "ExplorerProbe 构造函数",
                })
                score -= 5.0

        if re.search(r'add\s*\(\s*instrument\s*\)\s*;', code_text):
            issues_data.append({
                "type": "method_call",
                "severity": "error",
                "class": "ExplorerProbe",
                "description": "调用了不存在的 add() 方法",
                "score_impact": -3.0,
                "suggestion": "请使用 this.scientificInstruments.add(instrument);",
                "location": "ExplorerProbe.addInstrument()",
            })
            score -= 3.0

    # 6. 检查 SpaceControlCenter main 方法
    if "class SpaceControlCenter" in code_text:
        if 'Spaceship(' in code_text and 'new Spaceship' not in code_text and 'new CargoShuttle' not in code_text:
            issues_data.append({
                "type": "new_keyword",
                "severity": "error",
                "class": "SpaceControlCenter",
                "description": "创建对象时可能漏掉了 'new' 关键字",
                "score_impact": -5.0,
                "suggestion": "Java中创建对象必须使用 new 关键字，如 new Spaceship(...)",
                "location": "SpaceControlCenter.main()",
            })
            score -= 5.0

        if 'Collections.sort(fleet)' in code_text:
            issues_data.append({
                "type": "comparable",
                "severity": "error",
                "class": "SpaceControlCenter",
                "description": "Spaceship 未实现 Comparable 接口，无法直接使用 Collections.sort() 排序",
                "score_impact": -5.0,
                "suggestion": "请让 Spaceship 实现 Comparable<Spaceship> 接口并重写 compareTo() 方法",
                "location": "SpaceControlCenter.main() - Collections.sort",
            })
            score -= 5.0

    # 7. 检查 toString() 方法签名
    tostring_missing = re.findall(r'public\s+String\s+toString\s*\(\s*\)\s*;', code_text)
    if tostring_missing:
        issues_data.append({
            "type": "tostring_body",
            "severity": "error",
            "class": "通用",
            "description": "存在未实现的 toString() 方法（只有声明没有方法体）",
            "score_impact": -3.0,
            "suggestion": "请在 toString() 方法中添加 return 语句",
            "location": "类中的 toString()",
        })
        score -= 3.0

    score = max(0, score)

    issues = [
        {
            "description": i["description"],
            "severity": i["severity"],
            "suggestion": i["suggestion"],
            "location": i.get("location", ""),
        }
        for i in issues_data
    ]

    if not issues:
        issues.append({
            "description": "代码结构完整，未发现明显错误",
            "severity": "info",
            "suggestion": "请检查逻辑正确性，并确保符合题目要求",
            "location": "全局",
        })

    return {
        "score": round(score, 1),
        "total_score": 100.0,
        "issues": issues,
        "summary": f"共发现 {len(issues)} 个问题，综合得分 {round(score, 1)}/100",
    }


def _build_comments_xml(issues: List[dict]) -> str:
    """构建 word/comments.xml 的完整内容（标准 OOXML 格式）

    批注内容质量标准：
    1. 指出具体错误位置和问题
    2. 解释为什么这是问题
    3. 提供正确的修改方法
    4. 如果适用，给出正确的代码示例
    """
    from docx.oxml.ns import qn
    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    lines = [
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>',
        f'<w:comments xmlns:w="{W_NS}">',
    ]

    for idx, issue in enumerate(issues):
        cid = str(idx)
        sev = issue.get("severity", "warning")
        sev_label = {"error": "错误", "warning": "警告", "info": "提示"}.get(sev, "提示")
        desc = issue.get("description", "")
        suggestion = issue.get("suggestion", "")
        location = issue.get("location", "")
        author = "Kimi AI 批改助手"
        initials = "AI"
        date = datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")

        # 构建详细的批注内容
        comment_parts = []

        # 第一部分：错误类型标签
        comment_parts.append(f"[{sev_label}]")

        # 第二部分：问题描述（简洁）
        if location and location != "全局":
            comment_parts.append(f"位置：{location}")
        comment_parts.append(f"问题：{desc}")

        # 第三部分：修改建议（详细）
        if suggestion:
            comment_parts.append(f"建议：{suggestion}")

        # 第四部分：如果是代码错误，添加相关提示
        if "this" in suggestion.lower() or "this" in desc.lower():
            comment_parts.append("提示：使用 this 关键字区分成员变量和局部变量")
        if "static" in suggestion.lower() or "this" in suggestion.lower():
            comment_parts.append("注意：静态方法中不能使用 this 关键字")
        if "clone" in desc.lower():
            comment_parts.append("提示：深拷贝需要克隆所有可变对象")
        if "super" in suggestion.lower():
            comment_parts.append("提示：子类构造函数必须先调用父类构造函数")

        # 组合批注文本
        comment_text = " | ".join(comment_parts)
        sev_color = "C00000" if sev == "error" else ("E36C09" if sev == "warning" else "00B050")
        sugg_text = comment_parts[-1] if len(comment_parts) > 2 else suggestion

        # 构建 XML 结构
        comment_xml = (
            f"<w:p>"
            f"<w:pPr><w:pStyle w:val=\"a9\"/></w:pPr>"
            f"<w:r>"
            f"<w:rPr><w:rStyle w:val=\"a8\"/></w:rPr>"
            f"<w:annotationRef/>"
            f"</w:r>"
            f"<w:r>"
            f"<w:rPr><w:b/><w:color w:val=\"{sev_color}\"/></w:rPr>"
            f"<w:t xml:space=\"preserve\">{sev_label}：</w:t>"
            f"</w:r>"
            f"<w:r>"
            f"<w:t xml:space=\"preserve\">{_escape_xml(desc)}</w:t>"
            f"</w:r>"
            f"<w:r>"
            f"<w:br/>"
            f"</w:r>"
            f"<w:r>"
            f"<w:rPr><w:color w:val=\"0070C0\"/></w:rPr>"
            f"<w:t xml:space=\"preserve\">→ {_escape_xml(sugg_text)}</w:t>"
            f"</w:r>"
            f"</w:p>"
        )

        lines.append(
            f'<w:comment w:id="{cid}" '
            f'w:author="{_escape_xml(author)}" '
            f'w:date="{date}" '
            f'w:initials="{initials}">'
            f"{comment_xml}"
            f"</w:comment>"
        )

    lines.append("</w:comments>")
    return "\n".join(lines)


def _escape_xml(text: str) -> str:
    """简单 XML 字符转义"""
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


def _norm_ws(s: str) -> str:
    """去除空白 + 去除括号及其内容（用于中文章节种子匹配）"""
    if not s:
        return ""
    s = re.sub(r"\s+", "", s)
    s = re.sub(r"[（(][^）)]*[）)]", "", s)
    return s


def _location_hints(location: str) -> List[str]:
    """
    从 issue 的 location 字段生成可在正文中匹配的短语列表。
    - 去掉括号及内容（如「实验数据记录（3步）」→「实验数据记录」）
    - 用罗马数字/英文小节名映射中文章节种子
    - 同时保留原始 location（用于包含匹配）
    """
    if not (location or "").strip():
        return []
    s0 = (location or "").strip()
    if s0.lower() in ("全局", "全局。", "n/a", "无", "未知", "none", "null", ""):
        return []

    s = s0
    for pfx in ("位置:", "位置：", "Location:", "LOCATION:", "Section:", "章节:", "章节：", "在", "位于"):
        if s.startswith(pfx):
            s = s[len(pfx) :].strip()
        elif s.lower().startswith(pfx.lower()):
            s = s[len(pfx) :].strip()

    hints: List[str] = []

    # 原始 location（保留，用于原文精确匹配）
    if len(s) >= 2:
        hints.append(s)

    # 去掉括号后的纯文本（用于章节种子匹配）
    bracket_free = re.sub(r"[（(][^）)]*[）)]", "", s).strip()
    if len(bracket_free) >= 2 and bracket_free != s:
        hints.append(bracket_free)

    low = bracket_free.lower()

    roman_map = {
        "i.": ["一、", "一．", "一.", "实验原理", "原理与实验内容", "实验原理与实验内容", "实验内容"],
        "ii.": ["二、", "二．", "实验仪器", "仪器"],
        "iii.": ["三、", "三．", "实验数据", "数据记录", "实验数据记录", "实验记录"],
        "iv.": ["四、", "四．", "数据分析", "实验数据分析", "结论", "分析及结论", "实验分析"],
        "v.": ["五、", "五．", "思考题"],
        "vi.": ["六、", "六．"],
        "1.": ["一、", "实验原理", "实验原理与实验内容"],
        "2.": ["二、", "实验仪器"],
        "3.": ["三、", "实验数据", "实验数据记录"],
        "4.": ["四、", "数据分析", "实验数据分析"],
        "5.": ["五、", "思考题"],
    }
    for key, seeds in roman_map.items():
        if low.startswith(key) or low.replace(".", "") == key.replace(".", "") or f" {key}" in low:
            hints.extend(seeds)

    en_to_zh = [
        ("lab principles", ["一、", "实验原理", "实验原理与实验内容"]),
        ("principles and content", ["一、", "实验原理", "实验原理与实验内容"]),
        ("lab equipment", ["二、", "实验仪器", "仪器"]),
        ("equipment", ["二、", "实验仪器"]),
        ("lab data", ["三、", "实验数据", "实验数据记录"]),
        ("data recording", ["三、", "实验数据", "实验数据记录"]),
        ("data analysis", ["四、", "数据分析", "实验数据分析"]),
        ("analysis", ["四、", "数据分析"]),
        ("conclusion", ["四、", "结论", "数据分析"]),
        ("thinking questions", ["五、", "思考题"]),
        ("questions", ["五、", "思考题"]),
    ]
    for en, seeds in en_to_zh:
        if en in low:
            hints.extend(seeds)

    for frag in re.findall(r"[一二三四五六七八九十]+[、.．][^\s，,;；（()）]{1,30}", bracket_free):
        hints.append(frag.strip())

    out: List[str] = []
    seen: set = set()
    for h in hints:
        t = (h or "").strip()
        if len(t) < 2:
            continue
        k = _norm_ws(t).lower()
        if k not in seen:
            seen.add(k)
            out.append(t)
    return out


def _add_comment_refs_to_run(run_el, cid: str, w_ns: str) -> None:
    """
    向段落中注入批注范围标记。
    正确 OOXML 结构（所有批注标记均为 w:p 的直接子元素）：

        <w:p>
          <w:commentRangeStart w:id="X"/>
          <w:r><w:t>被标注的文本</w:t></w:r>
          <w:commentRangeEnd   w:id="X"/>
          <w:r><w:rPr><w:rStyle w:val="CommentReference"/></w:rPr>
               <w:commentReference w:id="X"/></w:r>
        </w:p>
    """
    # 找到段落 (w:p)
    parent = run_el.getparent()
    if parent is None:
        return

    if parent.tag == f"{{{w_ns}}}p":
        para = parent
    else:
        # run 可能被 w:hyperlink / w:del 等包裹
        para = parent.getparent()
        if para is None or para.tag != f"{{{w_ns}}}p":
            return

    # 检查该段落是否已有同 id 的 commentRangeStart（防止重复注入）
    existing = para.find(f"{{{w_ns}}}commentRangeStart[@{{{w_ns}}}id='{cid}']")
    if existing is not None:
        return  # 已注入过

    # 找到 run 在段落中的位置
    children = list(para)
    try:
        run_idx = children.index(run_el)
    except ValueError:
        return

    # 1. commentRangeStart → 插入到 run 之前
    commentRangeStart = OxmlElement("w:commentRangeStart")
    commentRangeStart.set(f"{{{w_ns}}}id", cid)
    para.insert(run_idx, commentRangeStart)

    # 2. commentRangeEnd → 插入到 run 之后（run_idx+1 因为刚才插入了一个元素）
    commentRangeEnd = OxmlElement("w:commentRangeEnd")
    commentRangeEnd.set(f"{{{w_ns}}}id", cid)
    para.insert(run_idx + 2, commentRangeEnd)

    # 3. commentReference run → 紧跟 commentRangeEnd
    ref_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(f"{{{w_ns}}}val", "CommentReference")
    rPr.append(rStyle)
    ref_run.append(rPr)
    commentRef = OxmlElement("w:commentReference")
    commentRef.set(f"{{{w_ns}}}id", cid)
    ref_run.append(commentRef)
    para.insert(run_idx + 3, ref_run)


def add_comments_to_docx(content: bytes, issues: List[dict], course: str = "Java OOP") -> bytes:
    """
    向 docx 文件注入 Word 批注（sidebar comment），批注锚点优先落在正文对应小节/关键词处。

    策略：
    1. 用 location 扩展出中英文章节别名，在正文段落中匹配（解决英文 location vs 中文标题）
    2. 再用 description/location/suggestion 中的代码类名等关键词匹配
    3. 仅对仍无法定位的 issue 在文末追加批注锚点（避免所有批注挤在「索引」段）
    """
    from docx.oxml import OxmlElement

    if not issues:
        return content

    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    w_ns = f"{{{W_NS}}}"
    CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
    REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"

    # ---- 1. 解析原始 ZIP ----
    try:
        original_zf = zipfile.ZipFile(io.BytesIO(content), "r")
    except zipfile.BadZipFile:
        return content

    # ---- 2. 读取 word/document.xml ----
    doc_xml_bytes = original_zf.read("word/document.xml")
    from lxml import etree
    tree = etree.fromstring(doc_xml_bytes)
    nsmap = {"w": W_NS}
    body = tree.find(".//w:body", nsmap)
    if body is None:
        return content

    # ---- 3. 为每个 issue 构建匹配关键词列表 ----
    issue_keywords = []
    import re as _re

    for idx, issue in enumerate(issues):
        desc = issue.get("description", "")
        location = issue.get("location", "")
        suggestion = issue.get("suggestion", "")

        # 从 description、location、suggestion 中提取关键词
        raw_tokens = (
            _re.findall(r'[A-Z][a-z]\w*', desc) +
            _re.findall(r'[A-Z]{2,}(?:\.[A-Z]+)*', desc) +
            _re.findall(r'[A-Z][a-z]+(?:\.[A-Z]+)*', desc) +
            _re.findall(r'[A-Z][a-z]\w*', location) +
            _re.findall(r'[A-Z]{2,}', location) +
            _re.findall(r'[A-Z][a-z]\w*', suggestion)
        )

        # 净化和去重
        clean_tokens = []
        seen_lower = set()
        for t in raw_tokens:
            if t.isascii() and len(t) >= 2:
                for part in t.split('.'):
                    pl = part.lower()
                    if pl not in seen_lower and len(part) >= 2:
                        seen_lower.add(pl)
                        clean_tokens.append(part)

        # 添加位置信息中的关键词
        if "EngineConfig" in location:
            clean_tokens.extend(["EngineConfig", "engineID", "constructor"])
        elif "Spaceship" in location:
            clean_tokens.extend(["Spaceship", "getTotalInstances", "clone"])
        elif "CargoShuttle" in location:
            clean_tokens.extend(["CargoShuttle", "cargoCapacity"])
        elif "ExplorerProbe" in location:
            clean_tokens.extend(["ExplorerProbe", "scientificInstruments", "ArrayList"])

        # 添加建议中的关键词
        if "this" in suggestion.lower():
            clean_tokens.append("this")
        if "static" in suggestion.lower():
            clean_tokens.append("static")
        if "super" in suggestion.lower():
            clean_tokens.append("super")
        if "clone" in suggestion.lower() or "clone" in desc.lower():
            clean_tokens.append("clone")
        if "Comparable" in suggestion or "Comparable" in desc:
            clean_tokens.append("Comparable")
        if "ArrayList" in suggestion or "ArrayList" in desc:
            clean_tokens.append("ArrayList")

        # 去重
        unique_kws = list(dict.fromkeys(clean_tokens))

        loc_hints = _location_hints(location)
        anchor = str(issue.get("anchor_snippet") or issue.get("anchor") or "").strip()

        # 额外增强：从 location 本身提取更多种子
        extra_hints = []
        for src in [location or "", anchor]:
            src = src.strip()
            if not src:
                continue
            for pfx in ("位置:", "位置：", "在", "位于", "步骤", "操作"):
                if src.startswith(pfx):
                    src = src[len(pfx):].strip()
            for m in re.finditer(r"第[一二三四五六七八九十\d]+[节步部分章节]", src):
                t = m.group().strip()
                if len(t) >= 3 and t not in extra_hints:
                    extra_hints.append(t)
            bare = re.sub(r"[（(][^）)]*[）)]", "", src).strip()
            if bare and bare != src and len(bare) >= 2:
                extra_hints.append(bare)
        loc_hints = extra_hints + loc_hints

        issue_keywords.append({
            "idx": idx,
            "cid": str(idx),
            "keywords": unique_kws,
            "location_hints": loc_hints,
            "anchor": anchor,
            "matched": False,
            "matched_para_idx": -1
        })

    # ---- 4. 遍历文档段落，匹配批注位置（仅原始正文段落，不含后续追加）----
    paragraphs = body.findall("w:p", nsmap)

    def _inject_on_first_text_run(para, para_idx: int, issue: dict) -> None:
        for run in para.findall("w:r", nsmap):
            t_el = run.find("w:t", nsmap)
            if t_el is not None and (t_el.text or "").strip():
                _add_comment_refs_to_run(run, issue["cid"], W_NS)
                issue["matched"] = True
                issue["matched_para_idx"] = para_idx
                return

    # 4a. 优先 anchor_snippet（模型从正文摘抄的短语），再按 location 别名匹配小节标题
    for para_idx, para in enumerate(paragraphs):
        para_text = ""
        for t_el in para.findall(".//w:t", nsmap):
            if t_el.text:
                para_text += t_el.text
        if not para_text.strip():
            continue
        norm_para = _norm_ws(para_text)

        for issue in issue_keywords:
            if issue["matched"]:
                continue
            a = issue.get("anchor") or ""
            if len(a) >= 2 and (a in para_text or _norm_ws(a) in norm_para):
                _inject_on_first_text_run(para, para_idx, issue)
                continue
            for hint in issue.get("location_hints") or []:
                nh = _norm_ws(hint)
                if len(nh) < 2:
                    continue
                if nh in norm_para or hint in para_text:
                    _inject_on_first_text_run(para, para_idx, issue)
                    break

    # 4b. 关键词匹配（代码标识符等）
    for para_idx, para in enumerate(paragraphs):
        para_text = ""
        for t_el in para.findall(".//w:t", nsmap):
            if t_el.text:
                para_text += t_el.text
        if not para_text.strip():
            continue
        para_text_lower = para_text.lower()

        for issue in issue_keywords:
            if issue["matched"]:
                continue
            for kw in issue["keywords"]:
                if kw.lower() in para_text_lower:
                    for run in para.findall("w:r", nsmap):
                        t_el = run.find("w:t", nsmap)
                        if t_el is None or not t_el.text:
                            continue
                        if kw.lower() in t_el.text.lower():
                            _add_comment_refs_to_run(run, issue["cid"], W_NS)
                            issue["matched"] = True
                            issue["matched_para_idx"] = para_idx
                            break
                    break

    unmatched = [i for i in issue_keywords if not i["matched"]]

    # ---- 5. 仅未能在正文定位的批注在文末落锚（避免重复 id、避免全部挤在文末）----
    if unmatched:
        empty_p = etree.SubElement(body, f"{w_ns}p")
        etree.SubElement(empty_p, f"{w_ns}r")

        title_p = etree.SubElement(body, f"{w_ns}p")
        title_r = etree.SubElement(title_p, f"{w_ns}r")
        title_rPr = etree.SubElement(title_r, f"{w_ns}rPr")
        etree.SubElement(title_rPr, f"{w_ns}b")
        title_t = etree.SubElement(title_r, f"{w_ns}t")
        title_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        title_t.text = "=== 以下批注未能在正文中自动定位，请手动对照 location 查看 ==="

        for issue in unmatched:
            idx = issue["idx"]
            issue_data = issues[idx]
            sev = issue_data.get("severity", "info")
            desc = issue_data.get("description", "")
            suggestion = issue_data.get("suggestion", "")
            location = issue_data.get("location", "")
            cid = issue["cid"]
            sev_icon = {"error": "[错误]", "warning": "[警告]", "info": "[提示]"}.get(sev, "[*]")

            p_el = etree.SubElement(body, f"{w_ns}p")
            commentRangeStart = etree.SubElement(p_el, f"{w_ns}commentRangeStart")
            commentRangeStart.set(f"{{{W_NS}}}id", cid)

            first_r = etree.SubElement(p_el, f"{w_ns}r")
            if location and location != "全局":
                content_text = f"{sev_icon} {desc} (位置: {location})"
            else:
                content_text = f"{sev_icon} {desc}"
            first_t = etree.SubElement(first_r, f"{w_ns}t")
            first_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            first_t.text = content_text

            commentRangeEnd = etree.SubElement(p_el, f"{w_ns}commentRangeEnd")
            commentRangeEnd.set(f"{{{W_NS}}}id", cid)

            ref_r = etree.SubElement(p_el, f"{w_ns}r")
            ref_rPr = etree.SubElement(ref_r, f"{w_ns}rPr")
            ref_rStyle = etree.SubElement(ref_rPr, f"{w_ns}rStyle")
            ref_rStyle.set(f"{{{W_NS}}}val", "CommentReference")
            commentRef = etree.SubElement(ref_r, f"{w_ns}commentReference")
            commentRef.set(f"{{{W_NS}}}id", cid)

            if suggestion:
                sugg_r = etree.SubElement(p_el, f"{w_ns}r")
                sugg_rPr = etree.SubElement(sugg_r, f"{w_ns}rPr")
                etree.SubElement(sugg_rPr, f"{w_ns}i")
                sugg_t = etree.SubElement(sugg_r, f"{w_ns}t")
                sugg_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                sugg_t.text = f"  → {suggestion}"

    # ---- 6. 文末一行说明（无批注锚点，避免侧栏再挂一条）----
    summary_p = etree.SubElement(body, f"{w_ns}p")
    summary_r = etree.SubElement(summary_p, f"{w_ns}r")
    summary_rPr = etree.SubElement(summary_r, f"{w_ns}rPr")
    etree.SubElement(summary_rPr, f"{w_ns}b")
    summary_t = etree.SubElement(summary_r, f"{w_ns}t")
    summary_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    if unmatched:
        summary_t.text = (
            f"=== 本次批改共 {len(issues)} 条批注：{len(issues) - len(unmatched)} 条已标注在正文，"
            f"{len(unmatched)} 条见上文「未定位」小节 ==="
        )
    else:
        summary_t.text = f"=== 本次批改共 {len(issues)} 条批注，均已标注在正文中 ==="

    new_doc_xml = etree.tostring(tree, xml_declaration=True, encoding="UTF-8", standalone=True)

    # ---- 6. 构建 word/comments.xml ----
    comments_xml_str = _build_comments_xml(issues)
    comments_xml_bytes = comments_xml_str.encode("utf-8")

    # ---- 7. 更新 [Content_Types].xml ----
    ct_xml_bytes = original_zf.read("[Content_Types].xml")
    ct_tree = etree.fromstring(ct_xml_bytes)
    CT = f"{{{CT_NS}}}"
    existing_override = ct_tree.find(f".//{{{CT_NS}}}Override[@PartName='/word/comments.xml']")
    if existing_override is None:
        override = etree.SubElement(ct_tree, f"{CT}Override")
        override.set("PartName", "/word/comments.xml")
        override.set(
            "ContentType",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
        )
    new_ct_xml = etree.tostring(ct_tree, xml_declaration=True, encoding="UTF-8", standalone=True)

    # ---- 8. 更新 word/_rels/document.xml.rels ----
    rels_xml_bytes = original_zf.read("word/_rels/document.xml.rels")
    rels_tree = etree.fromstring(rels_xml_bytes)
    REL = f"{{{REL_NS}}}"
    existing_rel = rels_tree.find(f".//{{{REL_NS}}}Relationship[@Target='comments.xml']")
    if existing_rel is None:
        new_rel = etree.SubElement(rels_tree, f"{REL}Relationship")
        max_id = 0
        for rel in rels_tree.findall(f"{{{REL_NS}}}Relationship"):
            rid = rel.get("Id", "")
            if rid.startswith("rId"):
                try:
                    max_id = max(max_id, int(rid[3:]))
                except ValueError:
                    pass
        new_rel.set("Id", f"rId{max_id + 1}")
        new_rel.set("Type", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments")
        new_rel.set("Target", "comments.xml")
    new_rels_xml = etree.tostring(rels_tree, xml_declaration=True, encoding="UTF-8", standalone=True)

    # ---- 9. 重新打包 ZIP ----
    output = io.BytesIO()
    with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as out_zf:
        for item in original_zf.infolist():
            if item.filename == "word/document.xml":
                out_zf.writestr(item, new_doc_xml)
            elif item.filename == "[Content_Types].xml":
                out_zf.writestr(item, new_ct_xml)
            elif item.filename == "word/_rels/document.xml.rels":
                out_zf.writestr(item, new_rels_xml)
            elif item.filename == "word/comments.xml":
                out_zf.writestr(item, comments_xml_bytes)
            else:
                out_zf.writestr(item, original_zf.read(item.filename))
        # 如果原 ZIP 里没有 comments.xml，追加进去
        if "word/comments.xml" not in original_zf.namelist():
            out_zf.writestr(
                zipfile.ZipInfo("word/comments.xml"),
                comments_xml_bytes
            )

    original_zf.close()
    output.seek(0)
    return output.read()


# ==========================
# API Endpoints
# ==========================

@router.post("/grade", response_model=ResponseModel)
async def grade_homework(
    homework_id: str = Form(...),
    student_id: str = Form(...),
    course: str = Form("Java OOP"),
    file: UploadFile = File(...),
):
    """
    AI批改作业文件
    """
    ensure_bucket()

    content = await file.read()
    original_filename = file.filename or "homework.docx"

    # 提取文本
    code_text = read_docx_text(content)
    if not code_text.strip():
        return ResponseModel(
            code=400,
            message="无法读取文档内容，请确保文件为docx格式",
            data=None
        )

    # AI评分
    grade_result = grade_java_code(code_text)
    issues = grade_result.get("issues", [])

    # 生成带批注的docx
    graded_content = add_comments_to_docx(content, issues, course)

    # 保存到MinIO
    review_id = f"review_{homework_id}_{datetime.now().strftime('%Y%m%d%H%M%S')}"
    graded_filename = original_filename.rsplit('.', 1)[0] + "_批改.docx"
    object_name = f"{student_id}/{review_id}/{graded_filename}"

    try:
        client = get_minio_client()
        client.put_object(
            BUCKET_REVIEW,
            object_name,
            io.BytesIO(graded_content),
            length=len(graded_content),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except S3Error:
        pass  # MinIO不可用时继续

    graded_file_url = f"http://{settings.MINIO_ENDPOINT}/{BUCKET_REVIEW}/{object_name}"

    # 保存元数据
    async with AsyncSessionLocal() as session:
        from sqlalchemy import select, update

        result = await session.execute(
            select(HomeworkReview).where(HomeworkReview.homework_id == homework_id)
        )
        existing = result.scalar_one_or_none()

        review_rec = None
        if existing:
            existing.review_id = review_id
            existing.original_filename = original_filename
            existing.graded_filename = graded_filename
            existing.graded_file_url = graded_file_url
            existing.file_size = len(graded_content)
            existing.score = grade_result.get("score")
            existing.total_score = grade_result.get("total_score")
            existing.status = "completed"
            existing.course = course
            existing.issue_count = len(issues)
            existing.issues_json = json.dumps(issues, ensure_ascii=False)
            existing.code_issues = issues
            existing.grading_model = "rule-based"
            existing.grading_time = datetime.utcnow()
            existing.original_content = code_text[:5000]
            existing.review_details = {"summary": grade_result.get("summary")}
            existing.updated_at = datetime.utcnow()
            review_rec = existing
        else:
            review_rec = HomeworkReview(
                review_id=review_id,
                homework_id=homework_id,
                student_id=student_id,
                original_filename=original_filename,
                graded_filename=graded_filename,
                graded_file_url=graded_file_url,
                file_size=len(graded_content),
                score=grade_result.get("score"),
                total_score=grade_result.get("total_score"),
                status="completed",
                course=course,
                issue_count=len(issues),
                issues_json=json.dumps(issues, ensure_ascii=False),
                code_issues=issues,
                grading_model="rule-based",
                grading_time=datetime.utcnow(),
                original_content=code_text[:5000],
                review_details={"summary": grade_result.get("summary")},
            )
            session.add(review_rec)

        await session.commit()
        await session.refresh(review_rec)

    return ResponseModel(
        code=200,
        message="AI批改完成",
        data={
            "review_id": review_rec.review_id,
            "homework_id": homework_id,
            "graded_filename": graded_filename,
            "graded_file_url": graded_file_url,
            "score": grade_result.get("score"),
            "total_score": grade_result.get("total_score"),
            "issue_count": len(issues),
            "issues": issues,
            "summary": grade_result.get("summary"),
        }
    )


@router.get("/list", response_model=ResponseModel)
async def list_reviews(
    student_id: str = Query(None),
    homework_id: str = Query(None),
    status: str = Query(None),
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=100),
):
    """列出已批改的作业"""
    async with AsyncSessionLocal() as session:
        from sqlalchemy import select, func
        query = select(HomeworkReview)
        count_q = select(func.count(HomeworkReview.id))

        conditions = []
        if student_id:
            conditions.append(HomeworkReview.student_id == student_id)
        if homework_id:
            conditions.append(HomeworkReview.homework_id == homework_id)
        if status:
            conditions.append(HomeworkReview.status == status)

        if conditions:
            query = query.where(*conditions)
            count_q = count_q.where(*conditions)

        total_res = await session.execute(count_q)
        total = total_res.scalar() or 0

        query = query.offset((page - 1) * page_size).limit(page_size).order_by(HomeworkReview.created_at.desc())
        res = await session.execute(query)
        reviews = res.scalars().all()

        items = []
        for r in reviews:
            issues = []
            try:
                if r.issues_json:
                    issues = json.loads(r.issues_json)
                elif r.code_issues:
                    issues = r.code_issues if isinstance(r.code_issues, list) else []
            except:
                pass

            items.append({
                "review_id": r.review_id,
                "homework_id": r.homework_id,
                "student_id": r.student_id,
                "original_filename": r.original_filename,
                "graded_filename": r.graded_filename,
                "graded_file_url": r.graded_file_url,
                "file_size": r.file_size,
                "score": r.score,
                "total_score": r.total_score,
                "status": r.status,
                "course": r.course,
                "issue_count": r.issue_count,
                "issues": issues,
                "created_at": r.created_at.isoformat() if r.created_at else None,
            })

        return ResponseModel(
            code=200,
            message="查询成功",
            data={"items": items, "total": total, "page": page, "page_size": page_size}
        )


@router.get("/{review_id}/download", response_model=ResponseModel)
async def download_review(review_id: str):
    """获取批改后文件的下载链接"""
    async with AsyncSessionLocal() as session:
        from sqlalchemy import select
        result = await session.execute(
            select(HomeworkReview).where(HomeworkReview.review_id == review_id)
        )
        review = result.scalar_one_or_none()

        if not review:
            raise HTTPException(status_code=404, detail="批改记录不存在")

        client = get_minio_client()
        object_name = f"{review.student_id}/{review.review_id}/{review.graded_filename}"

        try:
            presigned_url = client.presigned_get_object(
                BUCKET_REVIEW,
                object_name,
                expires=timedelta(seconds=3600)
            )
        except S3Error:
            presigned_url = review.graded_file_url

        return ResponseModel(
            code=200,
            message="获取下载链接成功",
            data={
                "download_url": presigned_url,
                "graded_filename": review.graded_filename,
                "expires_at": datetime.utcnow().isoformat(),
            }
        )


@router.get("/{review_id}", response_model=ResponseModel)
async def get_review(review_id: str):
    """获取批改详情"""
    async with AsyncSessionLocal() as session:
        from sqlalchemy import select
        result = await session.execute(
            select(HomeworkReview).where(HomeworkReview.review_id == review_id)
        )
        review = result.scalar_one_or_none()

        if not review:
            raise HTTPException(status_code=404, detail="批改记录不存在")

        issues = []
        try:
            if review.issues_json:
                issues = json.loads(review.issues_json)
            elif review.code_issues:
                issues = review.code_issues if isinstance(review.code_issues, list) else []
        except:
            pass

        return ResponseModel(
            code=200,
            message="查询成功",
            data={
                "review_id": review.review_id,
                "homework_id": review.homework_id,
                "student_id": review.student_id,
                "original_filename": review.original_filename,
                "graded_filename": review.graded_filename,
                "graded_file_url": review.graded_file_url,
                "file_size": review.file_size,
                "score": review.score,
                "total_score": review.total_score,
                "status": review.status,
                "course": review.course,
                "issue_count": review.issue_count,
                "issues": issues,
                "created_at": review.created_at.isoformat() if review.created_at else None,
            }
        )


@router.delete("/{review_id}", response_model=ResponseModel)
async def delete_review(review_id: str):
    """删除批改记录"""
    async with AsyncSessionLocal() as session:
        from sqlalchemy import select, delete
        result = await session.execute(
            select(HomeworkReview).where(HomeworkReview.review_id == review_id)
        )
        review = result.scalar_one_or_none()

        if not review:
            raise HTTPException(status_code=404, detail="批改记录不存在")

        # 删除 MinIO 文件
        client = get_minio_client()
        object_name = f"{review.student_id}/{review.review_id}/{review.graded_filename}"
        try:
            client.remove_object(BUCKET_REVIEW, object_name)
        except S3Error:
            pass

        await session.execute(
            delete(HomeworkReview).where(HomeworkReview.review_id == review_id)
        )
        await session.commit()

        return ResponseModel(code=200, message="删除成功")


# ============================================================
# Kimi AI 智能批改端点
# ============================================================

@router.get("/ai-grade/{homework_id}", response_model=ResponseModel)
async def ai_grade_homework(homework_id: str):
    """
    使用 Kimi AI 对作业进行智能批改：
    1. 从 MinIO 下载原始文件
    2. 调用 Kimi API 进行评分和批注生成
    3. 向 docx 文件注入 Word 批注（comment）
    4. 上传带批注的文件到 MinIO
    5. 保存批改记录，返回批改结果
    """
    ensure_bucket()

    # 1. 从数据库查询作业记录
    async with AsyncSessionLocal() as session:
        from sqlalchemy import select
        result = await session.execute(
            select(Homework).where(Homework.homework_id == homework_id)
        )
        homework_rec = result.scalar_one_or_none()
        if not homework_rec:
            raise HTTPException(status_code=404, detail="作业不存在")

    student_id = homework_rec.student_id
    original_filename = homework_rec.filename
    file_url = homework_rec.file_url

    # 2. 从 MinIO 下载原始文件
    client = get_minio_client()
    object_name = f"{student_id}/{homework_id}/{original_filename}"
    bucket_name = settings.MINIO_BUCKET_HOMEWORK

    try:
        response = client.get_object(bucket_name, object_name)
        file_content = response.read()
        response.close()
        response.release_conn()
    except S3Error:
        raise HTTPException(status_code=404, detail="文件不存在于 MinIO，请确认上传成功")

    # 3. 从 docx/doc 提取文本内容（旧版 .doc 会带 [DOC_FORMAT] 前缀，注入批注时用空白 docx）
    code_text = read_docx_text(file_content)
    if not code_text.strip():
        return ResponseModel(
            code=400,
            message="无法读取文档内容，请确保文件为 docx 格式",
            data=None
        )
    is_legacy_doc = code_text.startswith("[DOC_FORMAT]")
    if is_legacy_doc:
        code_text = code_text[len("[DOC_FORMAT]"):].strip()

    # 4. 调用 Kimi API 进行批改
    from common.integration.kimi import kimi_client

    # 获取课程类型，用于选择合适的批改提示词
    course = homework_rec.course or "通用"
    course_lower = course.lower()

    # 根据课程类型选择批改系统提示词
    if "java" in course_lower or "oop" in course_lower or "面向对象" in course:
        system_prompt = (
            "你是一位专业的计算机课程教师，擅长 Java 面向对象编程、算法设计的教学与批改。\n"
            "你的任务是仔细阅读学生提交的作业代码，给出专业、细致、有建设性的批注和修改建议。\n"
            "请严格按照以下 JSON 格式输出，不要包含任何其他文字：\n"
            "{\n"
            '  "score": <0-100的分数>,\n'
            '  "total_score": 100,\n'
            '  "issues": [\n'
            '    {\n'
            '      "description": "<具体问题描述>",\n'
            '      "severity": "<error|warning|info>",\n'
            '      "location": "<代码位置或类名>",\n'
            '      "anchor_snippet": "<从学生代码中逐字复制的连续片段8~40字符，用于在Word中定位批注；无法定位时填空字符串>",\n'
            '      "suggestion": "<修改建议>"\n'
            '    }\n'
            '  ]\n'
            '}\n'
            "重要：anchor_snippet 必须来自学生提交的代码原文（如某行代码、类名行），不要用概括性描述代替。\n"
            "请全面检查：\n"
            "- 类定义是否完整（public class、字段、方法）\n"
            "- 构造函数参数赋值是否正确使用 this 关键字\n"
            "- 静态方法中是否错误使用了 this 关键字\n"
            "- clone() 方法是否正确实现深拷贝\n"
            "- 父类构造函数调用是否使用 super()\n"
            "- toString() 等方法是否完整实现（有方法体）\n"
            "- 集合类型（ArrayList等）是否正确初始化\n"
            "- Comparable 接口是否正确实现\n"
            "- 其他代码规范和逻辑问题\n"
            "如果作业质量很高，给出 90-100 分；如果有严重错误，给出 60 分以下。"
        )
    elif "excel" in course_lower or "电子表格" in course or "数据处理" in course:
        system_prompt = (
            "你是一位专业的计算机教师，擅长 Microsoft Excel 操作与实验报告批改。\n"
            "你的任务是仔细批阅学生提交的 Excel 实验报告，给出专业、有针对性的批注和改进建议。\n"
            "请严格按照以下 JSON 格式输出，不要包含任何其他文字：\n"
            "{\n"
            '  "score": <0-100的分数>,\n'
            '  "total_score": 100,\n'
            '  "issues": [\n'
            '    {\n'
            '      "description": "<具体问题描述>",\n'
            '      "severity": "<error|warning|info>",\n'
            '      "location": "<操作步骤序号或位置>",\n'
            '      "anchor_snippet": "<从学生实验报告正文中逐字复制的连续片段8~40字，如章节标题「三、实验数据记录」或该节首句；用于Word定位>",\n'
            '      "suggestion": "<修改建议或正确做法>"\n'
            '    }\n'
            '  ]\n'
            '}\n'
            "重要：location 可用中文章节名；anchor_snippet 必须与报告原文完全一致（含标点），不要用英文改写标题。\n"
            "批改要点：\n"
            "- 检查操作步骤是否完整、准确\n"
            "- 检查公式是否正确（如 RANK 函数、COUNTIF/COUNTIFS、条件格式等）\n"
            "- 检查数据引用是否准确（列标、行号是否正确）\n"
            "- 检查是否按题目要求设置格式\n"
            "- 检查实验总结是否到位\n"
            "- 指出学生做得好的地方，给予肯定\n"
            "评分标准：\n"
            "- 优秀(90-100): 操作准确完整，步骤清晰\n"
            "- 良好(75-89): 基本完成，有小问题\n"
            "- 及格(60-74): 完成主要内容，但有明显错误\n"
            "- 不及格(<60): 未完成或错误较多"
        )
    elif "word" in course_lower or "文档" in course or "实验报告" in course or "报告" in course:
        system_prompt = (
            "你是一位专业的计算机课程教师，擅长各类办公软件操作和实验报告批改。\n"
            "你的任务是仔细批阅学生提交的实验报告或文档作业，给出专业、有建设性的批注和建议。\n"
            "请严格按照以下 JSON 格式输出，不要包含任何其他文字：\n"
            "{\n"
            '  "score": <0-100的分数>,\n'
            '  "total_score": 100,\n'
            '  "issues": [\n'
            '    {\n'
            '      "description": "<具体问题描述>",\n'
            '      "severity": "<error|warning|info>",\n'
            '      "location": "<具体位置或章节>",\n'
            '      "anchor_snippet": "<从学生报告正文中逐字复制的连续片段8~40字，如「一、实验原理与实验内容」；用于Word定位>",\n'
            '      "suggestion": "<改进建议>"\n'
            '    }\n'
            '  ]\n'
            '}\n'
            "重要：anchor_snippet 必须与报告原文完全一致，不要用英文标题（如 I. Lab Principles）代替中文章节标题。\n"
            "批改要点：\n"
            "- 内容是否完整（实验目的、原理、步骤、结果、结论）\n"
            "- 步骤描述是否准确、清晰\n"
            "- 截图或数据是否正确\n"
            "- 分析是否有深度\n"
            "- 格式是否规范\n"
            "- 鼓励创新和深入思考"
        )
    else:
        # 通用批改提示词
        system_prompt = (
            "你是一位专业、严格的教师，擅长各类作业和实验报告的批改。\n"
            "你的任务是仔细批阅学生提交的作业，给出专业、细致、有建设性的批注和建议。\n"
            "请严格按照以下 JSON 格式输出，不要包含任何其他文字：\n"
            "{\n"
            '  "score": <0-100的分数>,\n'
            '  "total_score": 100,\n'
            '  "issues": [\n'
            '    {\n'
            '      "description": "<具体问题描述>",\n'
            '      "severity": "<error|warning|info>",\n'
            '      "location": "<具体位置>",\n'
            '      "anchor_snippet": "<从学生作业正文中逐字复制的连续片段8~40字，用于在Word中定位批注>",\n'
            '      "suggestion": "<改进建议>"\n'
            '    }\n'
            '  ]\n'
            '}\n'
            "重要：每条批注尽量给出 anchor_snippet（原文摘抄），便于批注落在错误所在段落。\n"
            "批改要求：\n"
            "- 仔细阅读作业内容，找出优点和不足\n"
            "- 问题描述要具体，指出哪里做得不对\n"
            "- 建议要实用，能帮助学生改进\n"
            "- 适当给予鼓励，增强学生信心\n"
            "- 评分要客观公正"
        )

    user_prompt = f"请批改以下学生作业：\n\n{code_text}\n\n请给出分数和详细的批注，格式为 JSON。"

    kimi_result = await kimi_client.chat(
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ],
        temperature=0.3,
        max_tokens=4096
    )

    # 5. 解析 Kimi 返回结果
    issues = []
    score = 100.0
    total_score = 100.0
    summary = ""

    if "choices" in kimi_result and len(kimi_result["choices"]) > 0:
        raw_content = kimi_result["choices"][0]["message"]["content"]
        # 提取 JSON 部分（可能有 markdown 代码块包裹）
        json_str = raw_content.strip()
        if json_str.startswith("```"):
            lines = json_str.split("\n")
            json_str = "\n".join(lines[1:-1] if lines[-1] == "```" else lines[1:])
        try:
            parsed = json.loads(json_str)
            score = float(parsed.get("score", 100))
            total_score = float(parsed.get("total_score", 100))
            issues = parsed.get("issues", [])
            summary = f"共发现 {len(issues)} 个问题，综合得分 {score}/{total_score}"
        except json.JSONDecodeError:
            summary = raw_content[:200]
    else:
        error_msg = kimi_result.get("error", "Kimi API 调用失败")
        return ResponseModel(
            code=500,
            message=f"Kimi 批改失败：{error_msg}",
            data=None
        )

    # 6. 向 docx 注入批注（旧版 .doc 无法在原文件上批注，仅生成含总结的新 docx）
    docx_source = file_content if not is_legacy_doc else b""
    graded_content = add_comments_to_docx(docx_source, issues, homework_rec.course or "Java OOP")

    # 7. 上传带批注的文件到 MinIO
    review_id = f"kimi_review_{homework_id}_{datetime.now().strftime('%Y%m%d%H%M%S')}"
    graded_filename = original_filename.rsplit('.', 1)[0] + "_Kimi批改.docx"
    review_object_name = f"{student_id}/{review_id}/{graded_filename}"

    try:
        client.put_object(
            BUCKET_REVIEW,
            review_object_name,
            io.BytesIO(graded_content),
            length=len(graded_content),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except S3Error as e:
        pass  # MinIO 不可用时继续

    graded_file_url = f"http://{settings.MINIO_ENDPOINT}/{BUCKET_REVIEW}/{review_object_name}"

    # 8. 保存批改记录到数据库
    async with AsyncSessionLocal() as session:
        from sqlalchemy import select, update as sql_update

        # 查找是否已有该作业的批改记录
        existing_result = await session.execute(
            select(HomeworkReview).where(HomeworkReview.homework_id == homework_id)
        )
        existing = existing_result.scalar_one_or_none()

        if existing:
            existing.review_id = review_id
            existing.original_filename = original_filename
            existing.graded_filename = graded_filename
            existing.graded_file_url = graded_file_url
            existing.file_size = len(graded_content)
            existing.score = score
            existing.total_score = total_score
            existing.status = "completed"
            existing.course = homework_rec.course
            existing.issue_count = len(issues)
            existing.issues_json = json.dumps(issues, ensure_ascii=False)
            existing.code_issues = issues
            existing.grading_model = "kimi-ai"
            existing.grading_time = datetime.utcnow()
            existing.original_content = code_text[:5000]
            existing.review_details = {"summary": summary, "kimi_raw": kimi_result.get("choices", [{}])[0].get("message", {}).get("content", "")[:500]}
            existing.updated_at = datetime.utcnow()
            review_rec = existing
        else:
            review_rec = HomeworkReview(
                review_id=review_id,
                homework_id=homework_id,
                student_id=student_id,
                original_filename=original_filename,
                graded_filename=graded_filename,
                graded_file_url=graded_file_url,
                file_size=len(graded_content),
                score=score,
                total_score=total_score,
                status="completed",
                course=homework_rec.course,
                issue_count=len(issues),
                issues_json=json.dumps(issues, ensure_ascii=False),
                code_issues=issues,
                grading_model="kimi-ai",
                grading_time=datetime.utcnow(),
                original_content=code_text[:5000],
                review_details={"summary": summary, "kimi_raw": kimi_result.get("choices", [{}])[0].get("message", {}).get("content", "")[:500]},
            )
            session.add(review_rec)

        await session.execute(
            sql_update(Homework).where(Homework.homework_id == homework_id).values(
                status="reviewed",
                updated_at=datetime.utcnow(),
            )
        )

        await session.commit()
        await session.refresh(review_rec)

    return ResponseModel(
        code=200,
        message="Kimi 智能批改完成",
        data={
            "review_id": review_rec.review_id,
            "homework_id": homework_id,
            "student_id": student_id,
            "original_filename": original_filename,
            "graded_filename": graded_filename,
            "graded_file_url": graded_file_url,
            "score": score,
            "total_score": total_score,
            "issue_count": len(issues),
            "issues": issues,
            "summary": summary,
            "status": "completed",
        }
    )
